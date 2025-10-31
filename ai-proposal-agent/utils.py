from __future__ import annotations
import io
import logging
from pathlib import Path
from typing import Dict, List, Optional
from slugify import slugify as _slugify
from markdown import markdown as md_to_html
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import LETTER

logger = logging.getLogger(__name__)


def setup_logging(log_path: Path) -> None:
    log_path.parent.mkdir(parents=True, exist_ok=True)
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s %(levelname)s %(name)s: %(message)s",
        handlers=[logging.FileHandler(log_path, encoding="utf-8"), logging.StreamHandler()],
    )


def clean_transcript(text: str) -> str:
    """Lightly clean transcript text."""
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())


def render_markdown_to_html(md_text: str) -> str:
    css = """
    <style>
    body { font-family: -apple-system, Segoe UI, Roboto, Arial, sans-serif; line-height: 1.5; color: #111; }
    h1, h2, h3 { color: #0b3d91; }
    .cover { border-bottom: 2px solid #eee; margin-bottom: 16px; padding-bottom: 8px; }
    footer { font-size: 12px; color: #666; margin-top: 32px; }
    </style>
    """
    return css + md_to_html(md_text or "")


def slugify(title: str) -> str:
    return _slugify(title or "proposal")


def export_to_docx(sections: Dict[str, str], metadata: Dict[str, str], logo_path: Optional[Path]) -> bytes:
    doc = Document()
    # Cover Page
    if logo_path and logo_path.exists():
        try:
            doc.add_picture(str(logo_path), width=Inches(1.2))
        except Exception as e:
            logger.warning("Failed to add logo: %s", e)
    doc.add_heading(metadata.get("project_title", "Proposal"), 0)
    doc.add_paragraph(f"Company: {metadata.get('company_name','')}")
    doc.add_paragraph(f"Client: {metadata.get('client_name','')}")
    doc.add_page_break()

    for title, content in sections.items():
        doc.add_heading(title, level=1)
        for para in (content or "").split("\n\n"):
            doc.add_paragraph(para)
        doc.add_page_break()

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def export_to_pdf(html_or_sections: Dict[str, str] | str, metadata: Dict[str, str]) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    width, height = LETTER
    text = c.beginText(72, height - 72)
    header = f"{metadata.get('project_title','Proposal')} - {metadata.get('company_name','')} for {metadata.get('client_name','')}"
    text.textLine(header)
    text.textLine("" )
    if isinstance(html_or_sections, dict):
        for title, content in html_or_sections.items():
            text.textLine(f"## {title}")
            for ln in (content or "").splitlines():
                for chunk in [ln[i:i+90] for i in range(0, len(ln), 90)]:
                    text.textLine(chunk)
            text.textLine("")
            c.drawText(text)
            c.showPage()
            text = c.beginText(72, height - 72)
    else:
        for ln in (html_or_sections or "").splitlines():
            for chunk in [ln[i:i+90] for i in range(0, len(ln), 90)]:
                text.textLine(chunk)
    c.drawText(text)
    c.showPage()
    c.save()
    return buf.getvalue()
