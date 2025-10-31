from __future__ import annotations
import io
import logging
from pathlib import Path
from typing import Dict, Optional
from slugify import slugify as _slugify
from markdown import markdown as md_to_html
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import LETTER
try:
    import pdfkit  # type: ignore
except Exception:  # pragma: no cover
    pdfkit = None

logger = logging.getLogger(__name__)


def setup_logging(log_path: Path) -> None:
    log_path.parent.mkdir(parents=True, exist_ok=True)
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s %(levelname)s %(name)s: %(message)s",
        handlers=[logging.FileHandler(log_path, encoding="utf-8"), logging.StreamHandler()],
    )


def clean_transcript(text: str) -> str:
    table = {
        0x2018: ord("'"),  # ‘
        0x2019: ord("'"),  # ’
        0x201C: ord('"'),  # “
        0x201D: ord('"'),  # ”
        0x2013: ord('-'),   # –
        0x2014: ord('-'),   # —
        0x2026: ord('.'),   # …
    }
    norm = (text or "").translate(table).replace("…", "...")
    norm = "\n".join(line.strip() for line in norm.splitlines() if line.strip())
    try:
        return norm.encode("utf-8", errors="ignore").decode("utf-8")
    except Exception:
        return norm


def render_markdown_to_html(md_text: str) -> str:
    css = """
    <style>
    body { background:#111; }
    .doc-wrap { display:flex; justify-content:center; padding:24px; }
    .doc-page { background:#fff; width:794px; min-height:1123px; padding:48px 64px; box-shadow:0 4px 20px rgba(0,0,0,0.25); border-radius:6px; }
    .doc-page h1, .doc-page h2, .doc-page h3 { color:#0b3d91; font-family: Segoe UI, Roboto, Arial, sans-serif; }
    .doc-page p, .doc-page li { font: 15px/1.7 "Segoe UI", Roboto, Arial, sans-serif; color:#222; }
    .doc-page hr { border:0; border-top:1px solid #e5e7eb; margin:24px 0; }
    .doc-meta { color:#475569; font-size:13px; margin-bottom:12px; }
    </style>
    """
    html = md_to_html(md_text or "", extensions=["tables", "sane_lists", "fenced_code"])
    return css + f"<div class='doc-wrap'><div class='doc-page'>{html}</div></div>"


def slugify(title: str) -> str:
    return _slugify(title or "proposal")


def export_to_docx(sections: Dict[str, str], metadata: Dict[str, str], logo_path: Optional[Path]) -> bytes:
    doc = Document()
    # Cover Page
    if logo_path and logo_path.exists():
        try:
            doc.add_picture(str(logo_path), width=Inches(1.2))
        except Exception as e:
            logger.warning("Failed to add logo: %s", e)
    doc.add_heading(metadata.get("project_title", "Proposal"), 0)
    doc.add_paragraph(f"Company: {metadata.get('company_name','')}")
    doc.add_paragraph(f"Client: {metadata.get('client_name','')}")
    doc.add_paragraph(f"Tone: {metadata.get('brand_tone','')}")
    doc.add_page_break()

    for title, content in sections.items():
        doc.add_heading(title, level=1)
        for para in (content or "").split("\n\n"):
            doc.add_paragraph(para)
        doc.add_page_break()

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def export_html(md_text: str, logo_b64: Optional[str] = None, logo_ext: str = "png") -> bytes:
    html = render_markdown_to_html(md_text)
    if logo_b64:
        logo_img = f"<img src='data:image/{logo_ext};base64,{logo_b64}' style='height:60px;margin-bottom:16px'/>"
        html = html.replace("<div class='doc-page'>", f"<div class='doc-page'>{logo_img}", 1)
    return html.encode("utf-8")


def export_to_pdf(html_or_sections: Dict[str, str] | str, metadata: Dict[str, str], logo_path: Optional[Path] = None) -> bytes:
    # Prefer pdfkit (wkhtmltopdf) to preserve HTML/CSS exactly like preview
    if pdfkit is not None and isinstance(html_or_sections, (dict, str)):
        try:
            if isinstance(html_or_sections, dict):
                md = "\n\n".join([f"# {k}\n\n{v}" for k, v in html_or_sections.items()])
            else:
                md = html_or_sections
            # Build same HTML as preview
            import base64
            logo_b64 = None
            if logo_path and Path(logo_path).exists():
                logo_b64 = base64.b64encode(Path(logo_path).read_bytes()).decode("ascii")
            html_bytes = export_html(md, logo_b64)
            # Configure wkhtmltopdf path if provided via env var
            import os
            wkhtml = os.environ.get("WKHTMLTOPDF_PATH")
            cfg = pdfkit.configuration(wkhtmltopdf=wkhtml) if wkhtml else None  # type: ignore
            options = {
                "enable-local-file-access": None,
                "quiet": None,
                "page-size": "A4",
                "margin-top": "10mm",
                "margin-bottom": "12mm",
                "margin-left": "12mm",
                "margin-right": "12mm",
                "print-media-type": None,
                "encoding": "UTF-8",
            }
            pdf_bytes: bytes = pdfkit.from_string(html_bytes.decode("utf-8"), False, configuration=cfg, options=options)  # type: ignore
            return pdf_bytes
        except Exception:
            pass

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    width, height = LETTER
    text = c.beginText(72, height - 72)
    header = f"{metadata.get('project_title','Proposal')} - {metadata.get('company_name','')} for {metadata.get('client_name','')}"
    text.textLine(header)
    text.textLine("")

    def draw_footer() -> None:
        pg = c.getPageNumber()
        footer = f"Page {pg}"
        c.setFont("Helvetica", 9)
        c.drawRightString(width - 72, 36, footer)

    def draw_logo() -> None:
        if logo_path and Path(logo_path).exists():
            try:
                # draw at top-left corner
                c.drawImage(str(logo_path), 72, height - 72 - 50, width=50, height=50, preserveAspectRatio=True, mask='auto')
            except Exception:
                pass

    if isinstance(html_or_sections, dict):
        draw_logo()
        for title, content in html_or_sections.items():
            text.textLine(f"## {title}")
            for ln in (content or "").splitlines():
                for chunk in [ln[i:i+90] for i in range(0, len(ln), 90)]:
                    text.textLine(chunk)
            text.textLine("")
            c.drawText(text)
            draw_footer()
            c.showPage()
            text = c.beginText(72, height - 72)
            draw_logo()
    else:
        for ln in (html_or_sections or "").splitlines():
            for chunk in [ln[i:i+90] for i in range(0, len(ln), 90)]:
                text.textLine(chunk)
    c.drawText(text)
    draw_footer()
    c.showPage()
    c.save()
    return buf.getvalue()


def export_html_string_to_pdf(html: str) -> bytes:
    """Convert already-rendered HTML (same as preview) to PDF.
    Uses pdfkit/wkhtmltopdf if available; otherwise returns a minimal ReportLab PDF of plain text.
    """
    if pdfkit is not None:
        import os, tempfile
        wkhtml = os.environ.get("WKHTMLTOPDF_PATH")
        cfg = pdfkit.configuration(wkhtmltopdf=wkhtml) if wkhtml else None  # type: ignore
        options = {
            "enable-local-file-access": None,
            "quiet": None,
            "page-size": "A4",
            "margin-top": "10mm",
            "margin-bottom": "12mm",
            "margin-left": "12mm",
            "margin-right": "12mm",
            "print-media-type": None,
            "encoding": "UTF-8",
        }
        # Try in-memory first
        try:
            return pdfkit.from_string(html, False, configuration=cfg, options=options)  # type: ignore
        except Exception:
            # Then try writing to a temporary file
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".html", mode="w", encoding="utf-8") as f:
                    f.write(html)
                    temp_path = f.name
                try:
                    return pdfkit.from_file(temp_path, False, configuration=cfg, options=options)  # type: ignore
                finally:
                    try:
                        os.remove(temp_path)
                    except Exception:
                        pass
            except Exception:
                # Escalate to caller so UI can indicate wkhtmltopdf issue
                raise
    # Fallback plain text PDF
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    width, height = LETTER
    text = c.beginText(72, height - 72)
    for ln in html.splitlines():
        for chunk in [ln[i:i+100] for i in range(0, len(ln), 100)]:
            text.textLine(chunk)
    c.drawText(text)
    c.showPage()
    c.save()
    return buf.getvalue()
